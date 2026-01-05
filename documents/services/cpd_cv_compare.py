import re
import json
from io import BytesIO
import fitz  # PyMuPDF
from docx import Document
import unicodedata

EDU_HEADERS = {
    'education', 'academic qualification', 'academic background',
    'qualifications', 'scholastic record',
    'projects', 'academic projects', 'technical projects', 
    'project experience', 'key projects'
}

SECTION_HEADERS = {
    'experience', 'work history', 'employment', 'professional experience',
    'skills', 'technical skills', 'core competencies',
    'languages', 'certifications', 'summary', 'profile', 'achievements',
    'interests', 'references', 'declaration'
}

INSTITUTION_KEYWORDS = [
    'university', 'college', 'school', 'institute', 'academy',
    'polytechnic', 'vidyalaya', 'campus', 'secondary', 'higher secondary',
    'board', 'foundation', 'center', 'centre'
]

DEGREE_KEYWORDS = [
    'bachelor', 'master', 'phd', 'doctorate', 'diploma', 'associate',
    'b.sc', 'b.e', 'b.tech', 'm.sc', 'm.tech', 'mba', 'bba', 'bca', 'mca',
    'bbs', 'slc', 'see', 'hse', 'plus two', '+2', 'intermediate',
    'engineering', 'computer science', 'information technology'
]

def normalize_text(text):
    if not text:
        return ""
    text = unicodedata.normalize('NFKD', str(text)).encode('ASCII', 'ignore').decode('utf-8')
    return re.sub(r'\s+', ' ', text).strip().lower()

def extract_year(text):
    pattern = r'\b(19|20)\d{2}\s*(?:[-–to/]\s*(?:(19|20)\d{2}|present|now|current))?\b'
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(0) if match else ""

def calculate_similarity(text1, text2):
    set1 = set(normalize_text(text1).split())
    set2 = set(normalize_text(text2).split())
    if not set1 or not set2: return 0.0
    return len(set1.intersection(set2)) / len(set1.union(set2)) if set1.union(set2) else 0.0

def is_academic_degree(text):
    norm = normalize_text(text)
    return any(k in norm for k in DEGREE_KEYWORDS)

def is_institution(text):
    norm = normalize_text(text)
    return any(k in norm for k in INSTITUTION_KEYWORDS)

def extract_cpd_content(file_obj):
    try:
        if hasattr(file_obj, 'read'):
            file_obj.seek(0)
            doc = Document(BytesIO(file_obj.read()))
        else:
            doc = Document(file_obj)

        qualifications = []
        DEGREE_HEADER_KEYWORDS = ['degree', 'qualification', 'title', 'program', 'activit', 'course', 'particular']
        INSTITUTION_HEADER_KEYWORDS = ['institu', 'university', 'college', 'organization', 'place', 'location', 'where']
        YEAR_HEADER_KEYWORDS = ['period', 'year', 'duration', 'date', 'time']

        for table in doc.tables:
            headers = []
            header_row_idx = -1
            for i, row in enumerate(table.rows[:5]):
                row_text = [normalize_text(c.text) for c in row.cells]
                if any(any(k in cell for k in DEGREE_HEADER_KEYWORDS) for cell in row_text):
                    headers = row_text
                    header_row_idx = i
                    break
            
            if header_row_idx == -1: continue

            idx_deg = next((i for i, h in enumerate(headers) if any(k in h for k in DEGREE_HEADER_KEYWORDS)), None)
            idx_org = next((i for i, h in enumerate(headers) if any(k in h for k in INSTITUTION_HEADER_KEYWORDS)), None)
            idx_per = next((i for i, h in enumerate(headers) if any(k in h for k in YEAR_HEADER_KEYWORDS)), None)

            for row in table.rows[header_row_idx + 1:]:
                if idx_deg is None or len(row.cells) <= idx_deg: continue
                
                degree_text = row.cells[idx_deg].text.strip()
                if not degree_text: continue

                inst_text = ""
                if idx_org is not None and len(row.cells) > idx_org:
                    inst_text = row.cells[idx_org].text.strip()

                per_text = ""
                if idx_per is not None and len(row.cells) > idx_per:
                    per_text = row.cells[idx_per].text.strip()
                    clean_year = extract_year(per_text)
                    if clean_year: per_text = clean_year

                qualifications.append({
                    "degree": degree_text,
                    "institution": inst_text,
                    "period": per_text,
                    "source": "CPD"
                })

        return {"document_type": "CPD", "academic_qualifications": qualifications, "count": len(qualifications)}
    except Exception as e:
        return {"error": f"CPD Error: {str(e)}", "academic_qualifications": []}

# cv parse function
def parse_education_lines(lines):
    entries = []
    current_entry = {"degree": "", "institution": "", "period": ""}
    
    for line in lines:
        norm = normalize_text(line)
        if not norm: continue
        
        is_date = bool(extract_year(line))
        is_inst = is_institution(line)
        is_degree_key = is_academic_degree(line)

        prev_entry_done = bool(current_entry["degree"] and (current_entry["period"] or current_entry["institution"]))
        
        is_new_title = is_degree_key or (
            not is_date and not is_inst and len(line) < 120 and
            (not current_entry["degree"] or prev_entry_done)
        )

        if is_new_title:
            if current_entry["degree"]: 
                entries.append(current_entry)
                current_entry = {"degree": "", "institution": "", "period": ""}
            
            current_entry["degree"] = line.strip()
            date_in_line = extract_year(line)
            if date_in_line:
                current_entry["period"] = date_in_line

        elif is_inst:
            if not current_entry["institution"]: current_entry["institution"] = line.strip()
        
        elif is_date:
            if not current_entry["period"]: current_entry["period"] = extract_year(line)

    if current_entry["degree"]:
        entries.append(current_entry)
    return entries

def extract_cv_content(file_obj):
    try:
        if hasattr(file_obj, 'read'):
            file_obj.seek(0)
            doc = fitz.open(stream=file_obj.read(), filetype="pdf")
        else:
            doc = fitz.open(file_obj)

        education_lines = []
        in_target_section = False
        
        for page in doc:
            blocks = page.get_text("blocks")
            for b in blocks:
                text = b[4]
                lines = text.split('\n')
                
                for line in lines:
                    norm_line = normalize_text(line)
                    
                    if any(h in norm_line for h in EDU_HEADERS): 
                        in_target_section = True
                        continue 
                    
                    if in_target_section and norm_line in SECTION_HEADERS and not any(h in norm_line for h in EDU_HEADERS):
                        in_target_section = False
                        break 
                    
                    if in_target_section:
                        education_lines.append(line)

        qualifications = parse_education_lines(education_lines)

        return {"document_type": "CV", "academic_qualifications": qualifications, "count": len(qualifications)}
    except Exception as e:
        return {"error": f"CV Error: {str(e)}", "academic_qualifications": []}


def compare_academic_qualifications(cpd_data, cv_data):
    cpd_quals = cpd_data.get("academic_qualifications", [])
    cv_quals = cv_data.get("academic_qualifications", [])
    
    comparison_results = []
    matches_found = 0
    
    for cpd in cpd_quals:
        best_match = None
        highest_score = 0.0
        
        for cv in cv_quals:
            deg_score = calculate_similarity(cpd['degree'], cv['degree'])
            inst_score = calculate_similarity(cpd['institution'], cv['institution'])
            year_score = 1.0 if (cpd['period'] and cv['period'] and (cpd['period'] in cv['period'] or cv['period'] in cpd['period'])) else 0.0
            
            total_score = (deg_score * 0.6) + (inst_score * 0.3) + (year_score * 0.1)
            
            if total_score > highest_score:
                highest_score = total_score
                best_match = cv
        
        status = "MATCHED" if highest_score >= 0.45 else "NOT_MATCHED"
        if status == "MATCHED": matches_found += 1
            
        comparison_results.append({
            "cpd_entry": cpd,
            "matched_cv_qualification": best_match,
            "match_status": status,
            "confidence_score": round(highest_score, 2)
        })

    total_cpd = len(cpd_quals)
    total_cv = len(cv_quals)
    
    if total_cpd == 0: overall = "NO_DATA_TO_COMPARE"
    elif matches_found == total_cpd: overall = "FULLY_MATCHED"
    elif matches_found > 0: overall = "PARTIALLY_MATCHED"
    else: overall = "NOT_MATCHED"

    return {
        "overall_status": overall,
        "match_summary": f"{matches_found}/{total_cpd} items verified.",
        "details": comparison_results,
        "counts": {
            "total_cpd": total_cpd,
            "total_cv": total_cv,
            "matched": matches_found,
            "not_matched": total_cpd - matches_found
        }
    }