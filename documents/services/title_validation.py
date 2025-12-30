from typing import List, IO, Optional
from rest_framework.exceptions import ValidationError
import pdfplumber
import docx
import re

class TitleValidationService:
    def validate_and_extract_title(self, file_obj: IO[bytes]) -> str | None:
        file_name = file_obj.name.lower()
        file_obj.seek(0)
        
        if file_name.endswith(".docx"):
            return self.extract_from_docx(file_obj)
        elif file_name.endswith(".pdf"):
            return self.extract_from_pdf(file_obj)
        else:
            raise ValidationError("Unsupported file type")
    
    def _extract_title_from_career_episode_pattern(self, text: str) -> Optional[str]:
        """
        Extract title from pattern: "Career Episode [integer] ... [Title] ... [integer].1 Introduction"
        """
        # Case-insensitive search
        pattern = r'(?:career|carrer)\s+episode\s+(\d+)'
        match = re.search(pattern, text, re.IGNORECASE)
        
        if not match:
            return None
        
        episode_number = match.group(1)
        start_pos = match.end()
     
        # Pattern: Look for Introduction markers. The title ends before "Introduction"
        # Try patterns in order of specificity:
        # 1. "[episode_number].1. Introduction" or "[episode_number].1 Introduction" (with episode number)
        # 2. Just "Introduction" (if it appears soon after and is followed by section numbers)
        patterns_to_try = [
            rf'{re.escape(episode_number)}\.1\.\s+introduction',  # "3.1. Introduction" with dot after 1
            rf'{re.escape(episode_number)}\.1\s+introduction',  # "1.1 Introduction" with space
            rf'{re.escape(episode_number)}\.1\s*introduction',  # "1.1 Introduction" with optional space
            rf'{re.escape(episode_number)}\.1\.\s*introduction',  # "3.1. Introduction" with optional space
            rf'{re.escape(episode_number)}\.\s*1\s+introduction',  # "1. 1 Introduction" (with space after dot)
        ]
        
        intro_match = None
        for intro_pattern in patterns_to_try:
            intro_match = re.search(intro_pattern, text[start_pos:], re.IGNORECASE)
            if intro_match:
                break
        
        # If no numbered Introduction found, try to find just "Introduction" 
        # but make sure it's followed by section numbering like "[episode_number].2" or similar
        if not intro_match:
            # Look for "Introduction" followed by section numbers (e.g., "Introduction   1.2. Background")
            intro_simple_pattern = r'\bintroduction\b'
            simple_matches = list(re.finditer(intro_simple_pattern, text[start_pos:start_pos+500], re.IGNORECASE))
            
            for simple_match in simple_matches:
                # Check if after "Introduction" there's a section number pattern
                after_intro = text[start_pos + simple_match.end():start_pos + simple_match.end() + 100]
                # Look for pattern like "[episode_number].[digit]" after Introduction
                section_pattern = rf'{re.escape(episode_number)}\.\d+'
                if re.search(section_pattern, after_intro):
                    intro_match = simple_match
                    break
        
        if not intro_match:
            return None
        
        # Extract title between "Career Episode [number]" and the Introduction marker
        title_text = text[start_pos:start_pos + intro_match.start()].strip()
        
        # Clean up the title (remove extra whitespace, newlines)
        title_text = re.sub(r'\s+', ' ', title_text).strip()
        
        # Make sure we're not returning "Career Episode" itself or empty/very short text
        if not title_text or len(title_text) < 3:
            return None
        
        # Make sure we're not accidentally returning "Career Episode" or similar
        if title_text.lower().strip().startswith(('career episode', 'carrer episode')):
            return None
        
        return title_text
    
    def extract_from_docx(self, file_obj: IO[bytes]) -> str | None:
        try:
            document = docx.Document(file_obj)
        except Exception:
            raise ValidationError("Invalid DOCX file")

        # First, check for "Career Episode" pattern (initial case)
        # Extract text from first 50 paragraphs to search for the pattern
        full_text = " ".join([p.text for p in document.paragraphs[:50]])
        career_episode_title = self._extract_title_from_career_episode_pattern(full_text)
        if career_episode_title:
            return career_episode_title
        
        # Check if "Career Episode" was found but pattern didn't match
        # If so, we should skip "Career Episode" text in fallback methods
        has_career_episode = bool(re.search(r'(?:career|carrer)\s+episode\s+\d+', full_text, re.IGNORECASE))

        for paragraph in document.paragraphs[:50]:
            text = paragraph.text.strip()
            if not text:
                continue
            # Skip paragraphs that are just "Career Episode" if we detected it but pattern didn't match
            if has_career_episode and re.match(r'^(?:career|carrer)\s+episode\s+\d+$', text, re.IGNORECASE):
                continue
            if paragraph.style.name.lower() == "title" and text:
                return text

        candidates = []
        for paragraph in document.paragraphs[:20]:
            text = paragraph.text.strip()
            if not text:
                continue
            # Skip paragraphs that are just "Career Episode" if we detected it but pattern didn't match
            if has_career_episode and re.match(r'^(?:career|carrer)\s+episode\s+\d+$', text, re.IGNORECASE):
                continue
            
            max_size = 0
            for run in paragraph.runs:
                if run.font.size:
                    if run.font.size.pt > max_size:
                        max_size = run.font.size.pt
            
            if max_size > 15:
                # print(f"DEBUG: Found candidate with size {max_size}: {text}")
                candidates.append((max_size, text))
        
        if candidates:
            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]

        # for paragraph in document.paragraphs[:50]:
        #     if paragraph.style.name.lower().startswith("heading 1") and paragraph.text.strip():
        #         print(f"DEBUG: Found 'Heading 1' style: {paragraph.text.strip()}")
        #         return paragraph.text.strip()

        for paragraph in document.paragraphs[:50]:
            text = paragraph.text.strip()
            if not text:
                continue
            # Skip paragraphs that are just "Career Episode" if we detected it but pattern didn't match
            if has_career_episode and re.match(r'^(?:career|carrer)\s+episode\s+\d+$', text, re.IGNORECASE):
                continue
            # print(f"DEBUG: Fallback to first text: {text}")
            return text

        return None

    def extract_from_pdf(self, file_obj: IO[bytes]) -> str | None:
        try:
            with pdfplumber.open(file_obj) as pdf:
                if not pdf.pages:
                    return None
                
                # Extract text from first 3 pages to search for the pattern
                full_text = ""
                for page_num in range(min(3, len(pdf.pages))):
                    page = pdf.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + " "
                
                has_career_episode = False
                if full_text:
                    career_episode_title = self._extract_title_from_career_episode_pattern(full_text)
                    if career_episode_title:
                        return career_episode_title
                    
                    # Check if "Career Episode" was found but pattern didn't match
                    has_career_episode = bool(re.search(r'(?:career|carrer)\s+episode\s+\d+', full_text, re.IGNORECASE))
                
                for page_num in range(min(3, len(pdf.pages))):
                    page = pdf.pages[page_num]
                    words = page.extract_words(extra_attrs=["size"])
                    
                    if not words:
                        continue
                    
                    lines = self._group_words_by_line(words, tolerance=5.0)
                    
                    if not lines:
                        continue
                    
                    max_size = max(
                        max(float(w["size"]) for w in line["words"])
                        for line in lines[:10] if line["words"]
                    )
                    
                    title_lines = []
                    found_title_start = False
                    
                    for line in lines[:15]:
                        line_words = line["words"]
                        if not line_words:
                            if found_title_start:
                                break
                            continue
                        
                        line_max_size = max(float(w["size"]) for w in line_words)
                        line_text = " ".join(w["text"] for w in line_words).strip()
                        
                        if abs(line_max_size - max_size) <= 1.0:
                            title_lines.append(line_text)
                            found_title_start = True
                        elif found_title_start:
                            break
                    
                    if title_lines:
                        full_title = " ".join(title_lines).strip()
                        # print(f"DEBUG: Extracted PDF title from page {page_num + 1}: {full_title}")
                        return full_title
                
                for page_num in range(min(3, len(pdf.pages))):
                    page = pdf.pages[page_num]
                    first_text = page.extract_text()
                    if first_text and first_text.strip():
                        title = first_text.split("\n")[0].strip()
                        # Skip if it's just "Career Episode" and we detected it but pattern didn't match
                        if has_career_episode:
                            if re.match(r'^(?:career|carrer)\s+episode\s+\d+$', title, re.IGNORECASE):
                                # Try next line
                                lines = first_text.split("\n")
                                for line in lines[1:]:
                                    line = line.strip()
                                    if line and not re.match(r'^(?:career|carrer)\s+episode\s+\d+$', line, re.IGNORECASE):
                                        title = line
                                        break
                        # print(f"DEBUG: Fallback PDF title from page {page_num + 1}: {title}")
                        return title
        
        except Exception as e:
            # print(f"DEBUG: PDF Extraction Error: {e}")
            raise ValidationError("Invalid PDF file")
        
        return None

    def _group_words_by_line(self, words: List[dict], tolerance: float = 5.0) -> List[dict]:
        if not words:
            return []

        sorted_words = sorted(words, key=lambda w: (round(float(w["top"]), 1), float(w["x0"])))
        
        lines = []
        current_line = []
        current_top = None

        for word in sorted_words:
            top = float(word["top"])
            
            if current_top is None or abs(top - current_top) <= tolerance:
                current_line.append(word)
                current_top = top if current_top is None else (current_top + top) / 2
            else:
                if current_line:
                    lines.append({"top": current_top, "words": current_line})
                current_line = [word]
                current_top = top
        
        if current_line:
            lines.append({"top": current_top, "words": current_line})
        
        return lines