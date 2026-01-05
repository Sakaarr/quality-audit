from __future__ import annotations

import base64
import io
import re
import uuid
from datetime import datetime
from typing import List, Dict

from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

from documents.domain import DocumentImage, SectionNode, UnifiedDocument


class DocxParser:
    """Parses DOCX documents into the unified document model."""

    def parse(self, uploaded_file) -> UnifiedDocument:
        raw_bytes = uploaded_file.read()
        uploaded_file.seek(0)
        document = DocxDocument(io.BytesIO(raw_bytes))
        # First, extract all images and create a mapping by relationship ID
        images = self._extract_images(document)
        image_map = {img.metadata.get("relationship_id"): img for img in images}
        tables = self._extract_tables(document)
        table_map = {table["id"]: table for table in tables}

        images = self._extract_images(document)
        image_map = {img.metadata.get("relationship_id"): img for img in images}

        tables = self._extract_tables(document)
        table_map = {table["id"]: table for table in tables}
        metadata = self._extract_metadata(document)
        sections = self._build_sections(document)
        images = self._extract_images(document)
        
        paragraphs = []
        table_index = 1
        
        for element in document.element.body:
            if element.tag.endswith('p'):
                for paragraph in document.paragraphs:
                    if paragraph._element == element:
                        text_with_markers = self._get_paragraph_text_with_images(paragraph, image_map)
                        if text_with_markers.strip():
                            paragraphs.append(text_with_markers)
                        break
            
            elif element.tag.endswith('tbl'):
                table_id = f"table-{table_index}"
                paragraphs.append(f"<<TABLE>>")
                table_index += 1
        text_payload = {"full_text": "\n".join(paragraphs), "paragraphs": paragraphs}

        return UnifiedDocument(
            source_type="docx",
            metadata=metadata,
            sections=sections,
            text=text_payload,
            tables=tables,
            images=images,
            extras={
                "paragraph_count": len([p for p in paragraphs if not p.strip().startswith("<<TABLE:")]),
                "section_count": len(sections),
                "table_count": len(tables),
                "image_count": len(images),
            },
        )

    def _get_paragraph_text_with_images(self, paragraph, image_map: Dict[str, DocumentImage]) -> str:
        """
        Iterates through runs to extract text AND detect where images are located.
        Returns a string with '<<IMAGE>>' markers inserted at the correct positions.
        """
        text_parts = []
        
        # XML namespace for Word processing drawings
        namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        for run in paragraph.runs:
            # 1. Add the text content of the run
            if run.text:
                text_parts.append(run.text)
            
            # 2. Check the underlying XML of the run for drawing elements
            drawings = run.element.findall(f'.//{namespace}drawing')
            picts = run.element.findall(f'.//{namespace}pict')
            
            if drawings or picts:
                # Try to find the relationship ID for this image
                image_id = self._get_image_rId_from_run(run)
                
                if image_id and image_id in image_map:
                    # Use the actual image identifier from metadata
                    identifier = image_map[image_id].identifier
                    text_parts.append(f"\n<<IMAGE>>\n")
                else:
                    # Fallback if we can't find the specific image
                    text_parts.append("\n<<IMAGE>>\n")

        return "".join(text_parts).strip()
    
    
    def _get_image_rId_from_run(self, run) -> str | None:
        """
        Extract the relationship ID (rId) from a run's XML that contains an image.
        """
        # Namespaces used in DOCX XML
        namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'v': 'urn:schemas-microsoft-com:vml'
        }
        
        # Try to find the relationship ID in the blip element (modern images)
        blips = run.element.findall('.//a:blip', namespaces)
        for blip in blips:
            embed = blip.get(f"{{{namespaces['r']}}}embed")
            if embed:
                return embed
        
        # Try to find in VML imagedata (older format images)
        imagedatas = run.element.findall('.//v:imagedata', namespaces)
        for imagedata in imagedatas:
            rid = imagedata.get(f"{{{namespaces['r']}}}id")
            if rid:
                return rid
        
        return None
    def _extract_metadata(self, document: DocxDocument) -> dict:
        props = document.core_properties

        def iso(value: datetime | None) -> str | None:
            return value.isoformat() if isinstance(value, datetime) else None

        metadata = {
            "author": props.author,
            "created": iso(props.created),
            "modified": iso(props.modified),
            "last_modified_by": getattr(props, "last_modified_by", None),
            "category": props.category,
            "comments": props.comments,
            "keywords": props.keywords,
            "subject": props.subject,
            "title": props.title,
        }
        return {key: value for key, value in metadata.items() if value}

    def _build_sections(self, document: DocxDocument) -> List[SectionNode]:
        sections: List[SectionNode] = []
        stack: List[SectionNode] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = getattr(getattr(paragraph, "style", None), "name", "")
            if style_name and style_name.lower().startswith("heading"):
                level = self._heading_level(style_name)
                node = SectionNode(title=text, level=level)

                while stack and stack[-1].level >= level:
                    stack.pop()

                if stack:
                    stack[-1].children.append(node)
                else:
                    sections.append(node)

                stack.append(node)
            elif stack:
                stack[-1].paragraphs.append(text)

        return sections

    def _heading_level(self, style_name: str) -> int:
        match = re.search(r"(\d+)", style_name)
        return int(match.group(1)) if match else 1

    def _extract_tables(self, document: DocxDocument) -> list:
        tables = []
        nested_table_counter = {}
        for index, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_content = self._extract_cell_content_with_nested_tables(
                        cell, nested_table_counter
                    )
                    row_data.append(cell_content)
                rows.append(row_data)
            tables.append(
                {
                    "id": f"table-{index}",
                    "row_count": len(rows),
                    "column_count": len(rows[0]) if rows else 0,
                    "data": rows,
                }
            )
        return tables

    def _extract_cell_content_with_nested_tables(self, cell, nested_table_counter: dict) -> str:
        content_parts = []
        namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        for element in cell._element:
            if element.tag.endswith('p'):
                text = ''.join(node.text for node in element.iter() if node.text)
                if text.strip():
                    content_parts.append(text.strip())
            
            elif element.tag.endswith('tbl'):
                if 'nested' not in nested_table_counter:
                    nested_table_counter['nested'] = 1
                else:
                    nested_table_counter['nested'] += 1
                
                nested_id = f"table-nested-{nested_table_counter['nested']}"
                content_parts.append(f"<<TABLE>>")
        
        return ' '.join(content_parts) if content_parts else cell.text.strip()
    
    def _extract_images(self, document: DocxDocument) -> List[DocumentImage]:
        images: List[DocumentImage] = []

        for rel in document.part.rels.values():
            if rel.reltype != RT.IMAGE:
                continue

            image_bytes = rel.target_part.blob
            identifier = f"docx-image-{uuid.uuid4().hex}"
            width = height = 0
            mime_type = "image/unknown"

            try:
                with Image.open(io.BytesIO(image_bytes)) as image:
                    width, height = image.size
                    if image.format:
                        mime_type = f"image/{image.format.lower()}"
            except Exception:
                # Keep defaults when the image cannot be opened.
                pass

            encoded = base64.b64encode(image_bytes).decode("utf-8")
            images.append(
                DocumentImage(
                    identifier=identifier,
                    mime_type=mime_type,
                    width=width,
                    height=height,
                    data=encoded,
                    metadata={"relationship_id": rel.rId, "size_bytes": len(image_bytes)},
                )
            )

        return images